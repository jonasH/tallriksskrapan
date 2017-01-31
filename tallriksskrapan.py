# -*- coding: utf-8 -*-
import requests
import urllib.request
import io
import json
from helpers import utf8text

from lxml import html
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine
from docx import Document

week_number = 0
lastWeek = "0"


def parse_vecka():
    answer = requests.get('http://www.vecka.nu')
    root = html.fromstring(answer.text)
    for child in root.xpath('//time'):
        global week_number
        week_number = child.text
        lastWeek = "52" if week_number == "1" else str(int(week_number) - 1) 
    return 'Det är nu vecka %s' % week_number

def parse_kompassen():
    ret = "### KOMPASSEN ###" + "\n"
    answer = requests.get('http://www.restaurangkompassen.se/index.php?option=com_content&view=article&id=64&Itemid=66')
    root = html.fromstring(answer.text)
    friday_found = False
    for child in root.xpath('//div[@class="screen"]/div/div/div'):
        if friday_found and child.text:
            ret += child.text
        elif child.text and "fredag" in child.text.lower():
            friday_found = True
    return ret

def parse_teknikparken():
    ret = "### TEKNIKPARKEN ###" + "\n"
    answer = requests.get('http://www.restaurangteknikparken.se/index.php?option=com_content&view=article&id=46')
    root = html.fromstring(answer.text)
    friday_found = False
    for child in root.xpath('//div[@class="screen"]/div/div/div'):
        if friday_found and child.text:
            ret += child.text
        elif child.text and "fredag" in child.text.lower():
            friday_found = True
    return ret

def parse_gs():
    ret = "### Gourmetservice ###" + "\n"
    answer = requests.get('http://www.geflegourmetservice.se/lunch.php')
    root = html.fromstring(answer.text)

    for child in root.xpath('//div[@class="left_holder"]/p')[1:3]:
        ret += child.text_content()
    return ret

def parse_hemlingby():
    ret = "### HEMLINGBY ###" + "\n"
    answer = requests.get('http://www.gavle.se/Uppleva--gora/Idrott-motion-och-friluftsliv/Friluftsliv-och-motion/Hemlingby-friluftsomrade/Hemlingbystugan/Fika-och-ata/')
    root = html.fromstring(answer.text)
    for child in root.xpath('//a'):
        if child.text and "meny vecka" in child.text.lower() and week_number in child.text.lower():
            hemlingby_link='http://www.gavle.se' + child.get('href')
            break
    textAsArray = parse_pdf(hemlingby_link)
    ret += getFoodFromPDFArray(textAsArray)
    return ret
    
#Takes url to pdf file and returns text split on newline into array
def parse_pdf(pdf_url):

    remote_file = urllib.request.urlopen(pdf_url).read()
    memory_file = io.BytesIO(remote_file)
    parser = PDFParser(memory_file)
    doc = PDFDocument()
    parser.set_document(doc)
    #Warning sometimes, error in pdf?
    doc.set_parser(parser)
    doc.initialize('')
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    ret = []
    # Process each page contained in the document.
    for pageIdx, page in enumerate(doc.get_pages()):
        ret.append([])
        interpreter.process_page(page)
        layout = device.get_result()
        for idx, lt_obj in enumerate(layout):
            if isinstance(lt_obj, LTTextBox) or isinstance(lt_obj, LTTextLine):
                if len(lt_obj.get_text().strip()) > 0:
                    ret[pageIdx].append((lt_obj.get_text().splitlines()))
    return ret


def getFoodFromPDFArray(pdfArray):
    correctWeek = False
    ret = ""
    for page in pdfArray:
        for idx, line in enumerate(page):
            #Check if correct week
            if ("vecka " +  week_number + ":") in line[0]:
                correctWeek = True;
                continue;
            #If correct week and day is fredag
            elif correctWeek and "fredag" in line[0].lower():

                #If the line is bigger than 1 it contains the food after 'fredag'
                if len(line) > 1:
                    for x in range(1, len(line)):
                        ret += line[x] + "\n"
                #The line only contained 'fredag' so the food is in the line after 'fredag'
                else:
                    line = page[idx+1]
                    for x in range(0, len(line)):
                        ret += line[x] + "\n"
                return ret

    return "Oops something went wrong"


def parse_gustafsbro():
    ret = "### Gustafsbro ###" + "\n"
    answer = requests.get('http://www.gavlelunch.se/gustafsbro.asp')
    root = html.fromstring(answer.text)
    friday_found = False

    #Get friday from table
    for weekdayTable in root.xpath('//body/font/table/tr[1]/td[1]/div/table'):
        for day in weekdayTable.xpath('tr[1]/td[1]/font/strong'):
            if day.text and "fredag" in day.text.lower():
                friday_found = True
                break

    #If friday is found print food
    if friday_found:
           for food in weekdayTable.xpath('tr[2]/td[1]/font/ul/li'):
               ret += food.text.strip() + "\n"
    else:
           ret += "Oops something went wrong"
    return ret


def parse_sodersKalla():
    ret = "### Söders källa ###" + "\n"
    url = ""
    answer = requests.get('http://www.soderskalla.se/restaurangen/')
    root = html.fromstring(answer.text)

    #Get url for menu
    for child in root.xpath('//a'):
        if child.text and ("lunchmeny v" + week_number) in child.text.lower():
            url = child.get('href')
            break
        #Check if menu has been updated from the week before
        elif child.text and ("lunchmeny v" + lastWeek) in child.text.lower():
             ret += "Menyn har ännu inte blivit uppdaterad"

    #Fetch document
    answer = requests.get(url)
    memory_file = io.BytesIO(answer.content)
    doc = Document(memory_file)

    food = ""
    #Parse document and look for fredag, food is in the index after fredag
    for idx, para in enumerate(doc.paragraphs):
        if "fredag" in para.text.lower():
            food = doc.paragraphs[idx+1].text + "\n"

    if food:
        ret += food
    else:
        ret += "Oops something went wrong"
    return ret

def parse_koket():
    ret = "### Köket ###" + "\n"

    answer = requests.get('http://koketlunch.se/meny.html')
    root = html.fromstring(answer.text)
    friday_found = False
    food = ""

    for line in root.xpath('//p/span'):
       #Get friday from table
        if line.text and "fredag" in line.text.lower():
                friday_found = True

        if friday_found:
            if line.text.strip():
                #Fix encodings and remove '-' in the beginning of the different foods
                #Removed "fredag" printout to look more like the other printouts //Robert
                if "fredag" in utf8text(line.text).lower():
                    pass
                    #ret += utf8text(line.text) + "\n"
                elif "stängt" in utf8text(line.text).lower():
                    food += utf8text(line.text) + "\n"
                else:
                    food += utf8text(line.text)[1:] + "\n"
            else:
                break
    if food:
        ret += food
    else:
        ret += "Oops something went wrong"
    return ret

def parse_kryddan():
    ret = "### Kryddan ###" + "\n"
    answer = requests.get('http://www.kryddan35.se/hem/')
    root = html.fromstring(answer.text)
    friday_found = False
    food = ""
    for child in root.xpath('//div[@id="veckans"]'):
        lines = child.text_content().split("\n")
        for line in lines:
            if friday_found and child.text:
                food += line + "\n"
            elif line and "fredag" in line.lower():
                friday_found = True
    if food:
        ret += food
    else:
        ret += "Oops something went wrong"
    return ret

def get_json_encode():
    
    vecka = parse_vecka()
    teknikparken = parse_teknikparken()
    kompassen= parse_kompassen()
    hemlingby = parse_hemlingby()
    gs = parse_gs()
    gustafsbro= parse_gustafsbro()
    sodersKalla = parse_sodersKalla()
    koket = parse_koket()
    kryddan = parse_kryddan()
    
    return json.dumps({'vecka':week_number, 'teknikparken':teknikparken, 'kompassen':kompassen, 'hemlingby':hemlingby, 'gs':gs, 'gustafsbro':gustafsbro, 'sodersKalla':sodersKalla, 'koket':koket, 'kryddan':kryddan})

def main():
    parse_vecka()
    print("\n" +
    parse_teknikparken() + "\n" +
    parse_kompassen() + "\n" +
    parse_hemlingby() + "\n" +
    parse_gs()+ "\n" +
    parse_gustafsbro() + "\n" +
    parse_sodersKalla() + "\n" +
    parse_koket() + "\n" +
    parse_kryddan())
    
if __name__ == '__main__':
    main()
